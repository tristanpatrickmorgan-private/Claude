"""Regenerate Tax_Filing_Guide_2025.docx from the .md.

Parses a focused subset of markdown (h1-h4, tables, bold, italic, links, lists, paragraphs, blockquotes)
and produces a clean Word document. The .md is the canonical source — this is for stakeholders
who prefer Word.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD = '/home/user/Claude/Tristan/projects/uk-rental/Tax_Filing_Guide_2025.md'
OUT = '/home/user/Claude/Tristan/projects/uk-rental/Tax_Filing_Guide_2025.docx'

INLINE_RE = re.compile(
    r'(\*\*[^*]+\*\*|'           # bold **text**
    r'\*[^*]+\*|'                 # italic *text*
    r'`[^`]+`|'                   # code `text`
    r'\[[^\]]+\]\([^)]+\))'       # link [text](url)
)


def add_inline(paragraph, text):
    """Add text with inline markdown (bold, italic, code, links) parsed."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
        elif token.startswith('['):
            lm = re.match(r'\[([^\]]+)\]\(([^)]+)\)', token)
            if lm:
                txt, url = lm.group(1), lm.group(2)
                run = paragraph.add_run(txt)
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.font.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_md(md_path):
    """Yield ('type', content) tuples representing structural elements."""
    with open(md_path) as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip blank lines
        if not stripped:
            i += 1
            continue

        # Block quote
        if stripped.startswith('>'):
            block_lines = []
            while i < len(lines) and (lines[i].lstrip().startswith('>') or not lines[i].strip()):
                if lines[i].strip().startswith('>'):
                    block_lines.append(lines[i].lstrip().lstrip('>').strip())
                elif not lines[i].strip() and block_lines:
                    # blank inside quote - stop
                    break
                i += 1
            yield ('quote', '\n'.join(block_lines))
            continue

        # Headings
        if stripped.startswith('#'):
            m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if m:
                level = len(m.group(1))
                yield ('heading', (level, m.group(2)))
                i += 1
                continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            yield ('hr', None)
            i += 1
            continue

        # Tables — must have | and next line is separator
        if '|' in stripped and i + 1 < len(lines) and re.match(r'^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$', lines[i+1]):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            yield ('table', (header, rows))
            continue

        # Unordered list
        if re.match(r'^\s*[-*]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*[-*]\s+', lines[i]):
                items.append(re.sub(r'^\s*[-*]\s+', '', lines[i].rstrip('\n')))
                i += 1
            yield ('ul', items)
            continue

        # Ordered list
        if re.match(r'^\s*\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                items.append(re.sub(r'^\s*\d+\.\s+', '', lines[i].rstrip('\n')))
                i += 1
            yield ('ol', items)
            continue

        # Plain paragraph (could span multiple lines until blank)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '>', '|', '-', '*')) and not re.match(r'^\s*\d+\.', lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        yield ('para', ' '.join(para_lines))


def build_docx(md_path, out_path):
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for kind, content in parse_md(md_path):
        if kind == 'heading':
            level, text = content
            h = doc.add_heading(level=min(level, 4))
            # Strip markdown from heading text
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            h.add_run(clean)
        elif kind == 'para':
            p = doc.add_paragraph()
            add_inline(p, content)
        elif kind == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            # Strip inline markdown so it renders cleanly in the docx
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
            run = p.add_run(clean)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif kind == 'ul':
            for item in content:
                p = doc.add_paragraph(style='List Bullet')
                add_inline(p, item)
        elif kind == 'ol':
            for item in content:
                p = doc.add_paragraph(style='List Number')
                add_inline(p, item)
        elif kind == 'table':
            header, rows = content
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = 'Light Grid Accent 1'
            # Header
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(re.sub(r'\*\*([^*]+)\*\*', r'\1', h))
                run.bold = True
            # Rows
            for ri, row in enumerate(rows):
                for j, val in enumerate(row[:len(header)]):
                    cell = table.rows[ri + 1].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_inline(p, val)
        elif kind == 'hr':
            doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == '__main__':
    build_docx(MD, OUT)
